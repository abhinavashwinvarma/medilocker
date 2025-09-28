''' Using python-docx API.'''
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class Doctor:
    def __init__(self,doctorID ,docName , signature, clinic, clinicAddress, clinicLogoPath):
        self.doctorID = doctorID
        self.docName = docName
        self.signature = signature
        self.clinic=clinic
        self.clinicAddress = clinicAddress
        self.clinicLogoPath = clinicLogoPath

class Prescription(Doctor):

    def __init__(self, patient):
        self.patient= patient
        self.date= datetime.now()
        self.medicines=[]
        #self.address = fetch address of clinic from sql db

    def line():    
        line = document.add_heading(' ',0)
        run = line.runs[0]
        run.font.size=Pt(1)
        run.bold=False
        line.paragraph_format.line_spacing=0

    def add_med(self):
        name=input("Enter name of medicine: ")
        dosage= input("Enter dosage: ")
        frequency=input("Enter frequency: ")
        duration=input("Enter duration of course: ")
        notes=input("Enter any additional notes: ")
        self.medicines.append([name, dosage, frequency, duration, notes])
       
        
    def header(self, document, docName, qualifications, clinic, clinicAddress):
        document.add_picture(self.clinicLogoPath, width = Inches(1.2))
        table = document.add_table(rows=2, cols=2)
        table.autofit = False
        topLeft = table.cell(0,0).paragraphs[0] 
        topRight = table.cell(0,1).paragraphs[0]
        bottomLeft= table.cell(1,0).paragraphs[0]  
        bottomRight = table.cell(1,1).paragraphs[0]
        
        topLeft.add_run(clinic)
        bottomLeft.add_run(clinicAddress)
        topRight.add_run(f'Dr. {docName}')
        topRight.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        bottomRight.add_run(qualifications)
        bottomRight.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        line = document.add_heading(' ',0)
        run = line.runs[0]
        run.font.size=Pt(1)
        run.bold=False
        line.paragraph_format.line_spacing=0

    def body(self,document, patientName, patientAge, diagnosis, medicines, signaturePath):
        date = document.add_paragraph(f'Date: {datetime.now()}')
        date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        document.add_paragraph(f'Patient name: {patientName}\nPatient Age: {patientAge}')
        document.add_paragraph(f'Diagnosis: {diagnosis}')
        
        self.line()

        Col1 = document.add_paragraph('Medicine name')
        #col2, col3
        self.line()

        table = document.add_table(rows=1, cols=5)
        row = table.rows[0].cells
        row[0].text = 'Medicine Name'
        row[1].text = 'Dosage'
        row[2].text = 'Freqeuncy'
        row[3].text = 'Duration'
        row[4].text = 'Notes'

        for name, dosage, frequency, duration, notes in medicines:
            row = table.add_row().cells
            row[0].text = name
            row[1].text = dosage
            row[2].text = frequency
            row[3].text = duration
            row[4].text = notes
    

document = Document()
    
doctor1= Doctor('123456', 'Dr. Anishwar', "Signature.png", 'Atharnavee Ayurvedha', '#4/3, Vivekananda Nagar Main Road, Nesapakkam., Chennai', "Cliniclogo.png")
Prescription.header(document, 'Anishwar Balaji', 'B.A.M.S', 'Atharvanee Ayurvedha', '#4/3, Vivekananda Nagar Main Road, Nesapakkam., Chennai', 'Cliniclogo.png')
Prescription.body(document,'Anishwar', 17, 'Demodiagnosis', [['Nimbadi', 'Patolamundi','Frequency 1', 'Duration 1', 'Notes'],['medicine 2', 'heavy dose', 'very frequent', 'long duration', 'notesssssssssssssssssssssssssss' ]], 'Signature.png')

prescription=doctor1.Prescription(doctor1,'Abhinav', ['Nimbani', 'Sopanam'])
prescription.generatePrescription()

            
            


