from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

def line(document):    
    line = document.add_heading(' ',0)
    run = line.runs[0]
    run.font.size=Pt(1)
    run.bold=False
    line.paragraph_format.line_spacing=0

def header(document, docName, qualifications, clinic, clinicAddress, clinicLogoPath):
    document.add_picture(clinicLogoPath, width = Inches(1.2))
    table = document.add_table(rows=2, cols=2)
    table.autofit = False
    topLeft = table.cell(0,0).paragraphs[0] 
    topRight = table.cell(0,1).paragraphs[0]
    bottomLeft= table.cell(1,0).paragraphs[0]  
    bottomRight = table.cell(1,1).paragraphs[0]
 
    topLeft.add_run(clinic).font.name = 'IBM Plex Mono'
    bottomLeft.add_run(clinicAddress).font.name = 'IBM Plex Mono'
    topRight.add_run(f'Dr. {docName}').font.name = 'IBM Plex Mono'
    topRight.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    bottomRight.add_run(qualifications).font.name = 'IBM Plex Mono'
    bottomRight.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    line(document)

def body(document, docName, qualifications, patientName, patientAge, diagnosis, medicines, signaturePath):
    date = document.add_paragraph(f'Date: {datetime.now()}')
    for run in date.runs:
        run.font.name = 'IBM Plex Mono'
    date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    patientDetails = document.add_paragraph(f'Patient name: {patientName}\nPatient Age: {patientAge}')
    for run in patientDetails.runs:
        run.font.name = 'IBM Plex Mono'
    diagnosisPara = document.add_paragraph(f'Diagnosis: {diagnosis}')
    for run in diagnosisPara.runs:
        run.font.name = 'IBM Plex Mono'

    
    line(document)

    table = document.add_table(rows=1, cols=5, style = 'Table Grid')
    table.autofit = False
    row = table.rows[0].cells
    col1 = row[0].paragraphs[0].add_run('Medicine Name')
    col1.font.name = 'IBM Plex Mono'
    col1.font.bold = True
    col2 = row[1].paragraphs[0].add_run('Dosage')
    col2.font.name = 'IBM Plex Mono'
    col2.font.bold = True
    col3 = row[2].paragraphs[0].add_run('Frequency')
    col3.font.name= 'IBM Plex Mono'
    col3.font.bold = True
    col4 = row[3].paragraphs[0].add_run('Duration')
    col4.font.name = 'IBM Plex Mono'
    col4.font.bold = True
    col5 = row[4].paragraphs[0].add_run('Notes')
    col5.font.name = 'IBM Plex Mono'
    col5.font.bold = True
    
    line(document)

    for name, dosage, frequency, duration, notes in medicines:
        row = table.add_row().cells
        entry1 = row[0].paragraphs[0].add_run(name)
        entry1.font.name = 'IBM Plex Mono'
        entry2 = row[1].paragraphs[0].add_run(dosage)
        entry2.font.name = 'IBM Plex Mono'
        entry3 = row[2].paragraphs[0].add_run(frequency)
        entry3.font.name = 'IBM Plex Mono'
        entry4 = row[3].paragraphs[0].add_run(duration)
        entry4.font.name = 'IBM Plex Mono'
        entry5 = row[4].paragraphs[0].add_run(notes)
        entry5.font.name = 'IBM Plex Mono'
    
    section = document.sections[0]
    footer=section.footer
    footer_run = footer.paragraphs[0].add_run()
    footer_run.add_picture(signaturePath, width = Inches (1.2))
    footer.add_paragraph().add_run(f'Dr. {docName}\n'+ qualifications).font.name = 'IBM Plex Mono'
    line(document)

def create_prescription(doctorID, doctorname, qualifications, patientName, patientAge, patientDiagnosis, medicines, clinic, clinicAddress, clinicLogoPath, docSignaturePath):
    document = Document()
    header(document, doctorname, qualifications, clinic, clinicAddress, clinicLogoPath)
    body(document, doctorname, qualifications, patientName, patientAge,patientDiagnosis, medicines, docSignaturePath)

    document.save('DoctorFiles/' + doctorID + '/Prescription.docx')


#figure out the horizontal line thing ==> Title level  = 0
#figure out the spacing between the logos at the top 
#figure out tables
#do you want to do autofill tables?
