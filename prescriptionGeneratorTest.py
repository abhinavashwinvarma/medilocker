from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

def line():    
    line = document.add_heading(' ',0)
    run = line.runs[0]
    run.font.size=Pt(1)
    run.bold=False
    line.paragraph_format.line_spacing=0

def header(document, docName, qualifications, clinic, clinicAddress, clinicLogoPath):
    document.add_picture(clinicLogoPath, width = Inches(1.2))
    table = document.add_table(rows=2, cols=2)
    table.autofit = False
    topLeft = table.cell(0,0).paragraphs[0] 
    topRight = table.cell(0,1).paragraphs[0]
    bottomLeft= table.cell(1,0).paragraphs[0]  
    bottomRight = table.cell(1,1).paragraphs[0]
 
    topLeft.add_run(clinic)
    bottomLeft.add_run(clinicAddress)
    topRight.add_run(f'Dr. {docName}')
    topRight.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    bottomRight.add_run(qualifications)
    bottomRight.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    line()

def body(document, docName, qualifications, patientName, patientAge, diagnosis, medicines, signaturePath):
    date = document.add_paragraph(f'Date: {datetime.now()}')
    date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    document.add_paragraph(f'Patient name: {patientName}\nPatient Age: {patientAge}')
    document.add_paragraph(f'Diagnosis: {diagnosis}')

    
    line()

    table = document.add_table(rows=1, cols=5)
    row = table.rows[0].cells
    row[0].text = 'Medicine Name'
    row[1].text = 'Dosage'
    row[2].text = 'Freqeuncy'
    row[3].text = 'Duration'
    row[4].text = 'Notes'
    
    line()

    for name, dosage, frequency, duration, notes in medicines:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = dosage
        row[2].text = frequency
        row[3].text = duration
        row[4].text = notes

    
    section = document.sections[0]
    footer=section.footer
    footer_run = footer.paragraphs[0].add_run()
    footer_run.add_picture(signaturePath, width = Inches (1.2))
    footer.add_paragraph(f'Dr. {docName}\n'+ qualifications)
    line()

document = Document()

header(document, 'Anishwar Balaji', 'B.A.M.S', 'Atharvanee Ayurvedha', '#4/3, Vivekananda Nagar Main Road, Nesapakkam., Chennai', 'cliniclogo.png')
body(document,'Anishwar Balaji','B.A.M.S','Abhinav', 17, 'Demodiagnosis', [['Nimbadi', 'light dose','Frequency 1', 'Duration 1', 'Notes'],['medicine 2', 'heavy dose', 'very frequent', 'long duration', 'notesssssssssssssssssssssssssss' ]], 'signature.jpg')



document.save('Prescription.docx')
#figure out the horizontal line thing ==> Title level  = 0
#figure out the spacing between the logos at the top 
#figure out tables
#do you want to do autofill tables?
